import os
import hashlib
import pefile
import requests
import json
import time
import tkinter as tk
from tkinter import filedialog
from docx import Document

# Constants
VT_API_KEY = 'API_KEY'
HA_API_KEY = 'API_KEY'
VT_SCAN_URL = 'https://www.virustotal.com/api/v3/files'
VT_REPORT_URL = 'https://www.virustotal.com/api/v3/files/{}'
HA_SUBMIT_URL = 'https://www.hybrid-analysis.com/api/v2/quick-scan/file'
HA_REPORT_SUMMARY_URL = 'https://www.hybrid-analysis.com/api/v2/overview/{}'

# Helper functions
def get_file_hash(filepath):
    sha256_hash = hashlib.sha256()
    with open(filepath, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def submit_file_to_vt(filepath):
    headers = {"x-apikey": VT_API_KEY}
    with open(filepath, "rb") as f:
        files = {"file": f}
        response = requests.post(VT_SCAN_URL, headers=headers, files=files)
    if response.status_code != 200:
        print(f"Failed to submit file to VirusTotal: {response.status_code}")
        return None
    return response.json()

def get_vt_report(file_hash):
    headers = {"x-apikey": VT_API_KEY}
    response = requests.get(VT_REPORT_URL.format(file_hash), headers=headers)
    if response.status_code != 200:
        print(f"Failed to get VirusTotal report: {response.status_code}")
        return None
    return response.json()

def submit_file_to_ha(filepath):
    headers = {
        "api-key": HA_API_KEY,
        "user-agent": "Falcon Sandbox"
    }
    with open(filepath, "rb") as f:
        files = {"file": f}
        data = {"scan_type": "all"}
        response = requests.post(HA_SUBMIT_URL, headers=headers, files=files, data=data)
    if response.status_code != 200:
        print(f"Failed to submit file to Hybrid Analysis: {response.status_code}")
        return None
    return response.json()

def get_ha_report(sha256_hash):
    headers = {
        "api-key": HA_API_KEY,
        "user-agent": "Falcon Sandbox"
    }
    response = requests.get(HA_REPORT_SUMMARY_URL.format(sha256_hash), headers=headers)
    if response.status_code != 200:
        print(f"Failed to get Hybrid Analysis report: {response.status_code}")
        return None
    return response.json()

def static_analysis(filepath):
    pe = pefile.PE(filepath)
    info = {
        "Entry Point": hex(pe.OPTIONAL_HEADER.AddressOfEntryPoint),
        "Image Base": hex(pe.OPTIONAL_HEADER.ImageBase),
        "Compile Time": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime(pe.FILE_HEADER.TimeDateStamp)),
        "Sections": [section.Name.decode().strip() for section in pe.sections]
    }
    return info

def extract_iocs(filepath):
    iocs = {"IPs": [], "URLs": [], "Domains": [], "Hashes": []}
    with open(filepath, "rb") as f:
        data = f.read().decode(errors='ignore')
        import re
        iocs["IPs"] = re.findall(r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b', data)
        iocs["URLs"] = re.findall(r'(https?://\S+)', data)
        iocs["Domains"] = re.findall(r'\b(?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+[a-z]{2,6}\b', data)
    return iocs

def create_report(filepath, vt_report, ha_report, static_info, iocs):
    doc = Document()
    doc.add_heading('Malware Analysis Report', 0)
    
    doc.add_heading('File Information', level=1)
    doc.add_paragraph(f'File: {os.path.basename(filepath)}')
    doc.add_paragraph(f'SHA-256: {get_file_hash(filepath)}')

    doc.add_heading('VirusTotal Report', level=1)
    if vt_report:
        attributes = vt_report['data']['attributes']
        doc.add_paragraph(f"Malicious Votes: {attributes['last_analysis_stats'].get('malicious', 'N/A')}")
        doc.add_paragraph(f"Suspicious Votes: {attributes['last_analysis_stats'].get('suspicious', 'N/A')}")
        doc.add_paragraph(f"VT Link: {vt_report['data']['links']['self']}")
        
        # Adding more detailed information
        if 'last_analysis_results' in attributes:
            doc.add_heading('Detailed Analysis Results', level=2)
            for engine, result in attributes['last_analysis_results'].items():
                doc.add_paragraph(f"{engine}: {result['result']}")
    else:
        doc.add_paragraph("VT Report not available")

    doc.add_heading('Hybrid Analysis Report', level=1)
    if ha_report:
        doc.add_paragraph(f"SHA-256: {ha_report.get('sha256', 'N/A')}")
        doc.add_paragraph(f"Threat Score: {ha_report.get('threat_score', 'N/A')}")
        doc.add_paragraph(f"Verdict: {ha_report.get('verdict', 'N/A')}")
        doc.add_paragraph(f"Hybrid Analysis Report Link: https://www.hybrid-analysis.com/sample/{ha_report.get('sha256', 'N/A')}")
        
        # Adding more detailed information
        if 'analysis' in ha_report:
            doc.add_heading('Detailed Analysis Results', level=2)
            analysis = ha_report['analysis']
            if 'processes' in analysis:
                doc.add_paragraph(f"Processes: {json.dumps(analysis['processes'], indent=4)}")
            if 'network' in analysis:
                doc.add_paragraph(f"Network Activity: {json.dumps(analysis['network'], indent=4)}")
    else:
        doc.add_paragraph("HA Report not available")

    doc.add_heading('Static Analysis', level=1)
    for key, value in static_info.items():
        doc.add_paragraph(f'{key}: {value}')

    doc.add_heading('Indicators of Compromise (IOCs)', level=1)
    for key, values in iocs.items():
        doc.add_paragraph(f'{key}: {", ".join(values)}')

    report_path = os.path.splitext(filepath)[0] + "_report.docx"
    doc.save(report_path)
    return report_path

def main():
    # File selection
    root = tk.Tk()
    root.withdraw()
    filepath = filedialog.askopenfilename(title="Select an executable file",
                                          filetypes=[("Executable files", "*.exe")])
    if not filepath:
        print("No file selected.")
        return

    # VirusTotal scan
    print("Submitting file to VirusTotal...")
    vt_response = submit_file_to_vt(filepath)
    if vt_response:
        print("VirusTotal submission successful.")

    file_hash = get_file_hash(filepath)
    print(f"File hash: {file_hash}")

    print("Retrieving VirusTotal report...")
    vt_report = get_vt_report(file_hash)
    if vt_report:
        print("VirusTotal report retrieved.")

    # Static analysis
    print("Performing static analysis...")
    static_info = static_analysis(filepath)
    print("Static analysis completed.")
    
    # Extract IOCs
    print("Extracting IOCs...")
    iocs = extract_iocs(filepath)
    print("IOCs extraction completed.")
    
    # Create preliminary report
    print("Creating preliminary report...")
    report_path = create_report(filepath, vt_report, {}, static_info, iocs)
    print(f"Preliminary report created: {report_path}")

    # Hybrid Analysis scan
    print("Submitting file to Hybrid Analysis...")
    ha_response = submit_file_to_ha(filepath)
    if ha_response and 'sha256' in ha_response:
        sha256_hash = ha_response['sha256']
        print(f"SHA-256: {sha256_hash}")
        print("Waiting for Hybrid Analysis report...")
        while True:
            ha_report = get_ha_report(sha256_hash)
            if ha_report and ha_report.get('threat_score') is not None:
                break
            print("Hybrid Analysis report not ready, waiting for 30 seconds...")
            time.sleep(30)  # Wait for 30 seconds before checking again

        # Update report with Hybrid Analysis results
        if ha_report:
            print("Updating report with Hybrid Analysis results...")
            report_path = create_report(filepath, vt_report, ha_report, static_info, iocs)
            print(f"Final report updated: {report_path}")
        else:
            print("Failed to fetch Hybrid Analysis full report.")
    else:
        print("Failed to submit file to Hybrid Analysis.")

if __name__ == "__main__":
    main()
